"""
Resume Parser Service — PDF, DOCX extraction using PyMuPDF, pdfplumber, python-docx
"""
import os
import io
from pathlib import Path
from typing import Dict, Any, Optional
from backend.utils.logger import get_logger
from backend.utils.nlp_utils import (
    extract_skills, extract_email, extract_phone, extract_name,
    extract_experience_years, extract_tech_skills, extract_soft_skills
)

logger = get_logger(__name__)


def parse_resume(file_path: str, file_type: str) -> Dict[str, Any]:
    """
    Main entry point — parse resume and return structured data.

    Args:
        file_path: Absolute path to uploaded file
        file_type: 'pdf' or 'docx'

    Returns:
        dict with all extracted fields
    """
    try:
        if file_type.lower() == "pdf":
            raw_text = _parse_pdf(file_path)
        elif file_type.lower() in ("docx", "doc"):
            raw_text = _parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        if not raw_text or len(raw_text.strip()) < 50:
            raise ValueError("Could not extract meaningful text from resume")

        return _structure_resume(raw_text)

    except Exception as e:
        logger.error(f"Resume parsing failed: {e}")
        raise


def _parse_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF (primary) with pdfplumber fallback."""
    raw_text = ""

    # Primary: PyMuPDF (fitz)
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        for page in doc:
            raw_text += page.get_text("text") + "\n"
        doc.close()
        if raw_text.strip():
            logger.info(f"PDF parsed with PyMuPDF ({len(raw_text)} chars)")
            return raw_text
    except Exception as e:
        logger.warning(f"PyMuPDF failed: {e}, trying pdfplumber...")

    # Fallback: pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    raw_text += text + "\n"
        if raw_text.strip():
            logger.info(f"PDF parsed with pdfplumber ({len(raw_text)} chars)")
            return raw_text
    except Exception as e:
        logger.error(f"pdfplumber also failed: {e}")

    raise ValueError("Could not extract text from PDF")


def _parse_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = []

        # Main body paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Tables (often contain structured info)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        raw_text = "\n".join(paragraphs)
        logger.info(f"DOCX parsed ({len(raw_text)} chars)")
        return raw_text

    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        raise ValueError(f"Could not extract text from DOCX: {e}")


def _structure_resume(raw_text: str) -> Dict[str, Any]:
    """
    Convert raw resume text into structured data.
    """
    import re

    # Basic extraction
    name = extract_name(raw_text)
    email = extract_email(raw_text)
    phone = extract_phone(raw_text)
    skills = extract_skills(raw_text)
    experience_years = extract_experience_years(raw_text)

    # Education extraction
    education = _extract_education(raw_text)

    # Work experience extraction
    work_experience = _extract_work_experience(raw_text)

    # Projects
    projects = _extract_projects(raw_text)

    # Certifications
    certifications = _extract_certifications(raw_text)

    # Summary
    summary = _extract_summary(raw_text)

    return {
        "candidate_name": name,
        "email": email,
        "phone": phone,
        "raw_text": raw_text,
        "extracted_skills": skills,
        "experience_years": experience_years,
        "education": education,
        "certifications": certifications,
        "projects": projects,
        "work_experience": work_experience,
        "summary": summary,
        "languages": _extract_languages(raw_text),
    }


def _extract_education(text: str) -> list:
    """Extract education entries from text."""
    import re
    education = []
    lines = text.split('\n')
    edu_keywords = ['b.tech', 'b.e', 'm.tech', 'mca', 'bca', 'b.sc', 'm.sc',
                    'bachelor', 'master', 'phd', 'mba', 'b.com', 'm.com',
                    'diploma', 'university', 'college', 'institute', 'school']
    
    # Pre-compile boundary matching regexes for keywords
    edu_kws_regex = [re.compile(rf'\b{re.escape(kw)}\b') for kw in edu_keywords]
    
    for i, line in enumerate(lines):
        line_lower = line.lower()
        if any(rx.search(line_lower) for rx in edu_kws_regex):
            edu_entry = {"degree": line.strip(), "year": "", "institution": ""}
            # Look for year
            year_match = re.search(r'\b(19|20)\d{2}\b', line)
            if year_match:
                edu_entry["year"] = year_match.group()
            # Try next/nearby lines for institution
            if i + 1 < len(lines) and lines[i + 1].strip():
                next_line = lines[i + 1].strip()
                # Don't capture bullet points or headers as the institution
                if not next_line.startswith(('•', '-', '*')) and len(next_line) > 3 and not any(h in next_line.lower() for h in ['experience', 'skills', 'projects', 'certifications', 'summary']):
                    edu_entry["institution"] = next_line
            if edu_entry["degree"]:
                education.append(edu_entry)
    return education[:5]  # Cap at 5


def _extract_work_experience(text: str) -> list:
    """Extract work experience entries."""
    import re
    experiences = []
    lines = text.split('\n')

    # Date patterns for job entries
    date_pattern = re.compile(
        r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}'
        r'|\d{4}\s*[-–]\s*(?:\d{4}|Present|Current)',
        re.IGNORECASE
    )

    job_keywords = ['engineer', 'developer', 'analyst', 'manager', 'intern',
                    'consultant', 'specialist', 'lead', 'architect', 'designer',
                    'associate', 'coordinator', 'executive']

    for i, line in enumerate(lines):
        line_lower = line.lower()
        if any(kw in line_lower for kw in job_keywords) and len(line) > 5:
            exp = {
                "title": line.strip(),
                "company": "",
                "duration": "",
                "description": [],
            }
            # Look for date on same or adjacent line
            date_match = date_pattern.search(line)
            if date_match:
                exp["duration"] = date_match.group()
            elif i + 1 < len(lines):
                date_match = date_pattern.search(lines[i + 1])
                if date_match:
                    exp["duration"] = date_match.group()
                    exp["company"] = lines[i + 1].replace(exp["duration"], "").strip()

            # Collect bullet points
            j = i + 2
            while j < len(lines) and j < i + 8:
                desc_line = lines[j].strip()
                if desc_line and (desc_line.startswith('•') or desc_line.startswith('-')
                                  or desc_line.startswith('*') or len(desc_line) > 20):
                    exp["description"].append(desc_line.lstrip('•-* '))
                j += 1

            if exp["title"]:
                experiences.append(exp)

    return experiences[:10]


def _extract_projects(text: str) -> list:
    """Extract project entries from resume."""
    projects = []
    lines = text.split('\n')
    project_section = False
    current_project = None
    
    # Headers that indicate a new section is starting, meaning Projects section has ended
    section_headers = {
        'education', 'experience', 'skills', 'certifications', 'credentials',
        'languages', 'interests', 'hobbies', 'awards', 'publications',
        'summary', 'profile', 'objective'
    }

    for line in lines:
        line_strip = line.strip()
        line_lower = line_strip.lower()

        if 'project' in line_lower and len(line_strip) < 30:
            project_section = True
            continue

        if project_section:
            # If line matches a section header, terminate project parsing
            if any(hdr in line_lower for hdr in section_headers) and len(line_strip) < 25:
                project_section = False
                if current_project:
                    projects.append(current_project)
                    current_project = None
                continue

            if line_strip and not line_strip.startswith(('•', '-', '*')):
                if current_project:
                    projects.append(current_project)
                current_project = {
                    "name": line_strip,
                    "description": "",
                    "technologies": [],
                }
            elif current_project and line_strip:
                desc = line_strip.lstrip('•-* ')
                if current_project["description"]:
                    current_project["description"] += " " + desc
                else:
                    current_project["description"] = desc
                # Extract tech from description
                from backend.utils.nlp_utils import extract_tech_skills
                techs = extract_tech_skills(desc)
                current_project["technologies"].extend(techs)

        if len(projects) >= 10:
            break

    if current_project:
        projects.append(current_project)

    return projects


def _extract_certifications(text: str) -> list:
    """Extract certification entries from resume."""
    import re
    certifications = []
    lines = text.split('\n')
    cert_keywords = ['certif', 'aws', 'azure', 'google cloud', 'pmp', 'scrum',
                     'microsoft', 'oracle', 'cisco', 'comptia', 'coursera',
                     'udemy', 'credential', 'certified']
    cert_section = False

    for line in lines:
        line_lower = line.lower().strip()
        if 'certification' in line_lower or 'credential' in line_lower:
            cert_section = True
            continue
        if cert_section and any(kw in line_lower for kw in cert_keywords):
            cert = line.strip().lstrip('•-* ')
            if cert and len(cert) > 3:
                certifications.append(cert)
        elif any(kw in line_lower for kw in cert_keywords) and 'certified' in line_lower:
            cert = line.strip().lstrip('•-* ')
            if cert:
                certifications.append(cert)

    return list(set(certifications))[:10]


def _extract_summary(text: str) -> str:
    """Extract professional summary from resume."""
    lines = text.split('\n')
    summary_keywords = ['summary', 'objective', 'profile', 'about', 'overview']
    in_summary = False
    summary_lines = []

    for line in lines:
        line_lower = line.lower().strip()
        if any(kw in line_lower for kw in summary_keywords) and len(line) < 50:
            in_summary = True
            continue
        if in_summary:
            if line.strip() and len(line.strip()) > 20:
                summary_lines.append(line.strip())
            elif not line.strip() and summary_lines:
                break
        if len(summary_lines) >= 5:
            break

    return ' '.join(summary_lines) if summary_lines else ""


def _extract_languages(text: str) -> list:
    """Extract programming/human languages."""
    human_langs = ['english', 'hindi', 'tamil', 'telugu', 'marathi', 'bengali',
                   'gujarati', 'kannada', 'malayalam', 'punjabi', 'french',
                   'german', 'spanish', 'arabic', 'chinese', 'japanese']
    text_lower = text.lower()
    return [lang.capitalize() for lang in human_langs if lang in text_lower]
